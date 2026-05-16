"""
Generator dokumentasi MasPart -> Word (.docx).
Jalankan: python dokumentasi/_build_doc.py
Output:   dokumentasi/Dokumentasi-MasPart.docx
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = Path(__file__).resolve().parent
OUT  = BASE / "Dokumentasi-MasPart.docx"

# Mapping file gambar
IMG = {
    "login":           BASE / "login.png",
    "cari_pn":         BASE / "cari pn.png",
    "cari_name":       BASE / "cari part name.png",
    "cari_foto":       BASE / "cari by foto.png",
    "hasil_cari_foto": BASE / "hasil cari by foto.png",
    "gambar_part":     BASE / "gambar part.png",
}

# ── Style helpers ─────────────────────────────────────────────────────
MP_GREEN      = RGBColor(0x16, 0x7A, 0x3C)
MP_INK        = RGBColor(0x10, 0x18, 0x28)
MP_INK_50     = RGBColor(0x6B, 0x72, 0x80)
MP_BG         = RGBColor(0xF6, 0xF8, 0xFB)


def _shade(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd   = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_heading(doc, text, level=1, color=None):
    h = doc.add_heading("", level=level)
    run = h.add_run(text)
    run.bold = True
    if color is not None:
        run.font.color.rgb = color
    return h


def add_para(doc, text, bold=False, italic=False, color=None, size=11, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if align is not None:
        p.alignment = align
    return p


def add_bullets(doc, items):
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def add_numbered(doc, items):
    for it in items:
        doc.add_paragraph(it, style="List Number")


def add_image(doc, path: Path, width_inch=6.2, caption=None):
    if not path.exists():
        add_para(doc, f"[Gambar tidak ditemukan: {path.name}]", italic=True, color=MP_INK_50)
        return
    doc.add_picture(str(path), width=Inches(width_inch))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(f"Gambar: {caption}")
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = MP_INK_50


def add_info_box(doc, title, lines, fill="EAF7EE"):
    """Box hijau ringan untuk Tips / Catatan."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = True
    cell = tbl.rows[0].cells[0]
    _shade(cell, fill)
    p_title = cell.paragraphs[0]
    r = p_title.add_run(title)
    r.bold = True
    r.font.color.rgb = MP_GREEN
    r.font.size = Pt(11)
    for ln in lines:
        p = cell.add_paragraph()
        rr = p.add_run(f"• {ln}")
        rr.font.size = Pt(10.5)
        rr.font.color.rgb = MP_INK
    doc.add_paragraph()  # spacing


# ── Build document ────────────────────────────────────────────────────
doc = Document()

# Default style
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

# Margin
for section in doc.sections:
    section.left_margin   = Inches(0.9)
    section.right_margin  = Inches(0.9)
    section.top_margin    = Inches(0.8)
    section.bottom_margin = Inches(0.8)

# ── Cover ─────────────────────────────────────────────────────────────
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title_p.add_run("MasPart")
r.bold = True
r.font.size = Pt(36)
r.font.color.rgb = MP_GREEN

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("PART NUMBER FINDER")
r.bold = True
r.font.size = Pt(13)
r.font.color.rgb = MP_INK_50

doc.add_paragraph()
tag = doc.add_paragraph()
tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tag.add_run("Dokumentasi & Panduan Penggunaan Aplikasi")
r.font.size = Pt(14)
r.bold = True
r.font.color.rgb = MP_INK

desc = doc.add_paragraph()
desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = desc.add_run(
    "Database part Sinotruk · Shantui · Weichai, lengkap dengan foto SIMS,\n"
    "harga real-time, stok opname per user, dan tools batch."
)
r.font.size = Pt(11)
r.font.color.rgb = MP_INK_50

doc.add_paragraph()
add_image(doc, IMG["login"], width_inch=5.5, caption="Halaman Login MasPart")

doc.add_page_break()

# ── Daftar Isi (manual) ──────────────────────────────────────────────
add_heading(doc, "Daftar Isi", level=1, color=MP_GREEN)
toc_items = [
    "1. Pendahuluan",
    "2. Akses & Login",
    "3. Modul Search Part Number",
    "4. Modul Search Part Name",
    "5. Modul Cari by Foto",
    "6. Modul Bandingkan 2 Part",
    "7. Modul Batch Download",
    "8. Modul Stok Opname",
    "9. Tips Umum & Troubleshooting",
]
for it in toc_items:
    p = doc.add_paragraph(it)
    p.paragraph_format.left_indent = Inches(0.3)

doc.add_page_break()

# ── 1. Pendahuluan ───────────────────────────────────────────────────
add_heading(doc, "1. Pendahuluan", level=1, color=MP_GREEN)
add_para(
    doc,
    "MasPart adalah aplikasi web pencarian sparepart yang mengindeks puluhan file "
    "katalog Excel (Sinotruk, Shantui, Weichai) sekaligus terhubung ke SIMS untuk "
    "mengambil foto, harga, dan informasi part secara real-time. Aplikasi ini "
    "menggantikan proses manual membuka katalog satu per satu menjadi pencarian "
    "instan lewat satu halaman."
)
add_heading(doc, "Modul Utama yang Didokumentasikan", level=2)
add_bullets(doc, [
    "Search Part Number — pencocokan exact pada kolom Part Number.",
    "Search Part Name — pencarian substring (case-insensitive) pada nama part.",
    "Cari by Foto — pencarian visual menggunakan model DINOv2.",
    "Bandingkan 2 Part — analisis interchange (bentuk + nama + warna).",
    "Batch Download — pencarian banyak PN sekaligus + katalog Excel siap pakai.",
    "Stok Opname — sesi opname per user, simpan draft, riwayat, dan ekspor laporan.",
])

doc.add_page_break()

# ── 2. Akses & Login ─────────────────────────────────────────────────
add_heading(doc, "2. Akses & Login", level=1, color=MP_GREEN)
add_para(
    doc,
    "Aplikasi diakses melalui browser. Setiap user wajib login menggunakan username "
    "dan password yang terdaftar di sistem (Supabase Auth). Sesi tetap aktif selama "
    "75 menit setelah login terakhir."
)

add_heading(doc, "Langkah Login", level=2)
add_numbered(doc, [
    "Buka URL aplikasi MasPart di browser (Chrome/Edge disarankan).",
    "Masukkan Username pada kolom yang tersedia.",
    "Masukkan Password (klik ikon mata untuk menampilkan/menyembunyikan).",
    "Klik tombol Login → diarahkan ke halaman utama Part Number Finder.",
])

add_image(doc, IMG["login"], caption="Halaman login MasPart")

add_info_box(doc, "Catatan Penting", [
    "Jika belum punya akun, hubungi admin (link 'Hubungi admin' di bawah form login).",
    "Indikator hijau 'SUPABASE LIVE' menandakan koneksi ke database autentikasi aktif.",
    "Sesi otomatis logout setelah 75 menit tanpa aktivitas — login ulang untuk lanjut.",
])

doc.add_page_break()

# ── 3. Search Part Number ────────────────────────────────────────────
add_heading(doc, "3. Modul Search Part Number", level=1, color=MP_GREEN)
add_para(
    doc,
    "Tab pertama dan default. Digunakan untuk mencari part berdasarkan Part Number "
    "secara EXACT pada kolom B (Part Number) di seluruh file Excel terindex. Jika "
    "tidak ditemukan di database lokal, sistem otomatis fallback ke SIMS."
)

add_heading(doc, "Cara Pakai", level=2)
add_numbered(doc, [
    "Pilih tab '🔢 Search Part Number'.",
    "Ketik Part Number lengkap pada kolom input, misal: 080V10100-6088 atau WG1642821034/1.",
    "Klik tombol '🔍 Cari' (atau tekan Enter).",
    "Hasil ditampilkan dalam tabel: File, Part Number, Part Name, Quantity, Stok, Sheet, Excel Row.",
    "Klik baris hasil untuk melihat detail / foto part (gambar dari SIMS).",
])

add_image(doc, IMG["cari_pn"], caption="Hasil pencarian Part Number 080V10100-6088")

add_heading(doc, "Tips Pencarian Part Number", level=2)
add_bullets(doc, [
    "Gunakan part number LENGKAP untuk match exact (mis. WG1642821034).",
    "Format dengan garis miring didukung: WG1642821034/1.",
    "Jika hasil 0, sistem otomatis mencoba SIMS — pastikan koneksi internet aktif.",
    "Untuk pencarian fleksibel berdasarkan deskripsi, gunakan tab Search Part Name.",
    "Jika hanya punya foto, lompat ke tab Cari by Foto.",
])

add_heading(doc, "Penjelasan Kolom Hasil", level=2)
add_bullets(doc, [
    "File — nama katalog Excel sumber (mis. NX280 6X4 (LZZ1BLVF)).",
    "Part Number — PN yang cocok dengan input.",
    "Part Name — deskripsi part (mis. Fuel injector (Ks1.5)).",
    "Quantity — jumlah part pada unit (data dari katalog).",
    "Stok — stok terkini (jika modul Stok aktif untuk user).",
    "Sheet — nama sheet/section di file Excel.",
    "Excel Row — nomor baris pada file Excel asli.",
])

doc.add_page_break()

# ── 4. Search Part Name ──────────────────────────────────────────────
add_heading(doc, "4. Modul Search Part Name", level=1, color=MP_GREEN)
add_para(
    doc,
    "Digunakan jika Anda hanya tahu nama atau deskripsi part. Sistem mencari "
    "SUBSTRING (case-insensitive) pada kolom D (Part Name). Cocok untuk eksplorasi "
    "atau memetakan semua varian dari kata kunci tertentu."
)

add_heading(doc, "Cara Pakai", level=2)
add_numbered(doc, [
    "Pilih tab '📝 Search Part Name'.",
    "Ketik kata kunci pada kolom input, misal: injector, baut roda, kampas rem.",
    "Klik tombol '🔍 Cari'.",
    "Semua row yang mengandung kata kunci ditampilkan dalam tabel hasil.",
    "Gunakan scroll/filter untuk menjelajah hasil yang jumlahnya banyak.",
])

add_image(doc, IMG["cari_name"], caption="Hasil pencarian Part Name dengan kata kunci 'injector'")

add_heading(doc, "Tips Pencarian Part Name", level=2)
add_bullets(doc, [
    "Pencarian case-insensitive — 'INJECTOR' = 'injector' = 'Injector'.",
    "Pencarian substring — kata kunci 'fuel' menemukan 'fuel injector', 'fuel pump', dsb.",
    "Hasil terlalu banyak? Gunakan kata kunci lebih spesifik (mis. 'kampas rem depan').",
    "Untuk pencarian presisi, gunakan Search Part Number jika sudah tahu PN-nya.",
])

doc.add_page_break()

# ── 5. Cari by Foto ──────────────────────────────────────────────────
add_heading(doc, "5. Modul Cari by Foto", level=1, color=MP_GREEN)
add_para(
    doc,
    "Pencarian berbasis kemiripan visual menggunakan model DINOv2. Cocok ketika "
    "Anda hanya punya foto part fisik tanpa tahu nomor part. Sistem mencocokkan "
    "embedding foto Anda dengan ribuan foto SIMS yang sudah terindex dan menampilkan "
    "Part Number paling mirip beserta skor similarity."
)

add_heading(doc, "Cara Pakai", level=2)
add_numbered(doc, [
    "Pilih tab '🖼️ Cari by Foto'.",
    "Klik area 'Upload File' lalu pilih foto part (JPG/PNG). Anda juga dapat menggunakan kamera (fitur tambahan).",
    "Tunggu hingga preview foto muncul di sisi kanan (label 'Preview foto query').",
    "(Opsional) Klik 'Pengaturan lanjut' untuk mengatur threshold similarity / top-K.",
    "Klik tombol hijau '🔍 Cari Part Sekarang'.",
    "Sistem mengembalikan daftar PN ter-mirip beserta foto SIMS dan skor similarity (%).",
])

add_image(doc, IMG["cari_foto"], caption="Halaman upload foto query")

add_heading(doc, "Contoh Hasil", level=2)
add_para(
    doc,
    "Setelah klik Cari, akan tampil panel 'Hasil Pencarian' dengan badge 'Match kuat' "
    "atau 'Match lemah' beserta top similarity (%). Hasil ditampilkan grid 2 kolom: "
    "thumbnail foto SIMS, Part Number, Part Name, dan skor kemiripan."
)
add_image(doc, IMG["hasil_cari_foto"], caption="Hasil Cari by Foto — 'Match kuat' dengan top similarity 85.3%")

add_heading(doc, "Tips Foto untuk Hasil Terbaik", level=2)
add_bullets(doc, [
    "Foto tampak SAMPING dengan latar polos (putih/abu) memberi hasil paling akurat.",
    "Pencahayaan merata — hindari bayangan keras yang mengubah bentuk part.",
    "Part terlihat utuh, tidak terpotong oleh batas frame.",
    "Hindari pantulan/glare logam yang berlebihan.",
    "Resolusi minimal 300×300 px; format JPG atau PNG.",
])

add_info_box(doc, "Indikator Similarity", [
    "≥ 80% — Match kuat, kemungkinan besar PN benar.",
    "55–80% — Match sedang, verifikasi manual disarankan.",
    "< 55% — Match lemah, coba foto sudut lain atau gunakan Search Part Name.",
])

doc.add_page_break()

# ── 6. Bandingkan 2 Part ─────────────────────────────────────────────
add_heading(doc, "6. Modul Bandingkan 2 Part", level=1, color=MP_GREEN)
add_para(
    doc,
    "Tools 'Interchange Analyzer' untuk mengecek apakah 2 Part Number bisa saling "
    "substitusi (interchange). Sistem mengambil foto dari SIMS untuk kedua PN, lalu "
    "menganalisis 3 sinyal:"
)
add_bullets(doc, [
    "BENTUK (utama, bobot 60–70%) — embedding visual DINOv2.",
    "NAMA (penguat, bobot 25%) — kemiripan teks Part Name dari SIMS.",
    "WARNA (info, bobot 15%) — palet warna dominan untuk konfirmasi.",
])

add_heading(doc, "Cara Pakai", level=2)
add_numbered(doc, [
    "Pilih tab '🔍 Bandingkan 2 Part'.",
    "Isi Part Number #1 (mis. WG1642821034).",
    "Isi Part Number #2 (mis. WG1642821035).",
    "Klik tombol '🔬 Cek Interchange'.",
    "Tunggu sistem mengambil foto + analisis (biasanya 2–5 detik).",
    "Sistem menampilkan kartu verdict besar (hijau / kuning / merah) + 3 skor detail.",
    "Scroll ke bawah untuk melihat pasangan foto terbaik (gambar #i dari PN1 vs #j dari PN2).",
])

add_heading(doc, "Interpretasi Verdict", level=2)
add_bullets(doc, [
    "Verdict KUAT (✓) — shape ≥ 75% atau name ≥ 85% → kemungkinan besar interchange.",
    "Verdict RAGU (?) — skor menengah → cek manual sebelum dipasang.",
    "Verdict TIDAK COCOK — bentuk berbeda signifikan → bukan interchange.",
    "Overall score = (Bentuk × 0.6) + (Nama × 0.25) + (Warna × 0.15).",
])

add_info_box(doc, "Catatan", [
    "Fitur ini membutuhkan SIMS aktif — kedua PN harus punya foto di SIMS.",
    "Jika salah satu PN tidak ditemukan di SIMS, akan muncul pesan error spesifik.",
    "Untuk part dengan warna tipikal (mis. hitam), skor warna kurang informatif.",
])

doc.add_page_break()

# ── 7. Batch Download ────────────────────────────────────────────────
add_heading(doc, "7. Modul Batch Download", level=1, color=MP_GREEN)
add_para(
    doc,
    "Pencarian massal: upload daftar Part Number lalu sistem akan mencari semua PN "
    "sekaligus dan menghasilkan KATALOG Excel siap pakai (lengkap dengan foto SIMS "
    "dan harga). Cocok untuk kebutuhan RFQ, quotation, atau audit stok."
)

add_heading(doc, "Cara Pakai", level=2)
add_numbered(doc, [
    "Pilih tab '📥 Batch Download'.",
    "Klik 'Download Template Input' untuk mendapatkan template Excel (opsional).",
    "Pilih metode input: 'Upload File Excel' atau 'Ketik Manual'.",
    "Upload: file .xlsx/.xls/.xlsm/.csv dengan daftar PN di kolom A.",
    "Manual: ketik PN satu per baris di text area.",
    "Sistem otomatis membuang duplikat (dipertahankan kemunculan pertama).",
    "Klik 'Preview Part Number' untuk memastikan daftar benar.",
    "Klik tombol '🔍 Proses Batch Search'.",
    "Tunggu progress bar selesai (pencarian lokal → fetch SIMS untuk yang tidak ditemukan → ambil gambar).",
    "Download hasil dalam dua bentuk: Excel ringkasan + katalog Excel berisi foto.",
])

add_heading(doc, "Format File Input", level=2)
add_bullets(doc, [
    "Kolom A berisi Part Number (1 PN per baris).",
    "Header opsional (mis. 'Part Number'); sistem otomatis mendeteksi.",
    "Format ekstensi yang didukung: .xlsx, .xls, .xlsm, .csv.",
    "Tidak ada batas keras jumlah PN, namun >500 PN bisa memakan waktu beberapa menit.",
])

add_heading(doc, "Output Batch Download", level=2)
add_bullets(doc, [
    "Ringkasan hasil pada layar: Total PN, Ditemukan, Tidak Ditemukan.",
    "Preview tabel: Part Number, Hasil (file), Sheet, Part Name, Qty, Stok, Status.",
    "Tombol download Katalog Excel — file gabungan berisi foto SIMS per PN.",
    "Untuk PN yang tidak ada di lokal, Part Name otomatis diambil dari SIMS.",
])

add_info_box(doc, "Tips Batch Download", [
    "Bersihkan duplikat di file input untuk mempercepat proses.",
    "Pastikan koneksi internet stabil — proses fetch SIMS berjalan paralel.",
    "Jika prosesnya panjang, JANGAN tutup tab browser sebelum selesai.",
    "Simpan file output ke folder lokal Anda — link download bersifat sementara.",
])

doc.add_page_break()

# ── 8. Stok Opname ───────────────────────────────────────────────────
add_heading(doc, "8. Modul Stok Opname", level=1, color=MP_GREEN)
add_para(
    doc,
    "Modul untuk menjalankan sesi STOK OPNAME per user. Setiap user punya sesi "
    "sendiri (tidak saling mengganggu) dan draft otomatis tersimpan di Supabase "
    "(atau file lokal sebagai fallback). Setelah difinalisasi, sesi masuk ke "
    "Riwayat Opname dan bisa diekspor sebagai laporan Excel."
)

add_heading(doc, "Alur Sesi Opname", level=2)
add_numbered(doc, [
    "Buka tab '📋 Stok Opname'.",
    "Langkah 1 — (Opsional) Klik 'Download Template Stok Awal' untuk template Excel.",
    "Langkah 2 — Siapkan file Excel dengan kolom: Part Number, Qty Sistem (atau Stok), opsional Part Name.",
    "Upload file pada 'Upload Excel data stok awal'.",
    "Periksa preview & statistik (Total PN, Qty kosong, Qty 0, Punya Part Name).",
    "Klik '✅ Buat Sesi Opname dari Data Ini' → sesi draft dibuat.",
    "Isi Qty Fisik per PN — manual (ketik) atau upload Excel hasil hitung.",
    "Lihat panel 'Selisih' untuk memantau PN yang berselisih dengan sistem.",
    "Jika sudah lengkap, klik 'Finalisasi Sesi' → sesi masuk Riwayat Opname.",
    "Pada Riwayat, klik '⬇️ Excel' untuk download laporan, atau '🗑️' untuk hapus.",
])

add_heading(doc, "Format File Stok Awal", level=2)
add_bullets(doc, [
    "Header bebas posisi — sistem otomatis mendeteksi kolom 'Part Number', 'Qty Sistem'/'Stok', 'Part Name'.",
    "Jika tanpa header: A=PN, B=Qty, C=Part Name.",
    "Format file: .xlsx, .xls, .xlsm.",
    "Qty bisa kosong/N/A — akan dianggap belum ada data sistem (selisih = qty fisik).",
])

add_heading(doc, "Format File Upload Qty Fisik", level=2)
add_bullets(doc, [
    "Kolom Part Number harus cocok dengan PN yang ada pada draft.",
    "Kolom Qty Fisik (atau 'Qty Hitung', 'Qty Actual') berisi hasil hitung lapangan.",
    "PN yang tidak ada di draft akan diabaikan dengan warning.",
    "Sistem otomatis menghitung Selisih = Qty Fisik − Qty Sistem.",
])

add_heading(doc, "Kolom Riwayat Opname", level=2)
add_bullets(doc, [
    "Session ID — kode unik 8 karakter (untuk traceability).",
    "Difinalisasi — timestamp finalisasi.",
    "Total PN — jumlah PN dalam sesi.",
    "Sudah Hitung — PN yang sudah diisi Qty Fisik.",
    "Cocok — PN dengan Qty Fisik = Qty Sistem.",
    "Berselisih — PN dengan selisih ≠ 0.",
    "Selisih Net — total selisih (positif/negatif).",
])

add_info_box(doc, "Catatan Backend", [
    "Badge '☁️ Supabase' di atas judul = sesi tersimpan online (multi-device).",
    "Badge '💾 File lokal' = sesi hanya tersimpan di server tempat aplikasi berjalan.",
    "Jika muncul error 'Tabel opname_sessions belum ada', hubungi admin untuk setup Supabase.",
    "Draft bertahan selama belum difinalisasi — boleh ditutup & lanjut kapan saja.",
])

doc.add_page_break()

# ── 9. Tips Umum & Troubleshooting ───────────────────────────────────
add_heading(doc, "9. Tips Umum & Troubleshooting", level=1, color=MP_GREEN)

add_heading(doc, "Tips Umum", level=2)
add_bullets(doc, [
    "Gunakan Search Part Number jika tahu PN (paling akurat & cepat).",
    "Gunakan Search Part Name untuk eksplorasi varian/deskripsi.",
    "Gunakan Cari by Foto sebagai fallback jika hanya punya foto fisik.",
    "Gunakan Bandingkan 2 Part SEBELUM order substitusi untuk meminimalkan risiko.",
    "Gunakan Batch Download untuk RFQ / quotation banyak PN sekaligus.",
    "Selalu mulai sesi opname dengan upload data stok terbaru, bukan data lama.",
])

add_heading(doc, "Detail Gambar Part (Quick Look)", level=2)
add_para(
    doc,
    "Pada hasil pencarian (PN/Name/Foto), klik baris → bagian 'Gambar Part' "
    "memperlihatkan foto SIMS dengan kontrol Prev/Next, Zoom In/Out, dan tombol "
    "Refresh dari SIMS (memaksa ambil ulang dari server)."
)
add_image(doc, IMG["gambar_part"], caption="Panel Gambar Part dengan kontrol Prev/Next & Zoom")

add_heading(doc, "Troubleshooting", level=2)
add_bullets(doc, [
    "Login gagal — pastikan username/password benar; cek indikator 'SUPABASE LIVE'.",
    "Hasil Search PN kosong — coba tab Search Part Name, atau biarkan SIMS fallback berjalan.",
    "Cari by Foto lambat — foto besar membutuhkan waktu inference; resize ke ≤ 1024px.",
    "Bandingkan 2 Part error 'tidak ada gambar SIMS' — PN tidak terdaftar di SIMS; coba alternatif PN.",
    "Batch Download stuck — periksa koneksi internet & jangan tutup tab.",
    "Sesi opname tidak tersimpan — cek badge backend; hubungi admin jika Supabase off.",
    "Auto-logout 75 menit — login ulang untuk melanjutkan pekerjaan.",
])

add_info_box(doc, "Hubungi Admin Jika", [
    "Membutuhkan akun baru atau reset password.",
    "Permission menu tidak sesuai (mis. tab tertentu tidak muncul).",
    "Error 'Modul tidak ditemukan' atau gangguan SIMS berkepanjangan.",
    "Setup tabel opname_sessions di Supabase belum dilakukan.",
])

# ── Penutup ──────────────────────────────────────────────────────────
doc.add_paragraph()
end = doc.add_paragraph()
end.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = end.add_run("— Selesai —")
r.bold = True
r.font.color.rgb = MP_GREEN
r.font.size = Pt(12)

ftr = doc.add_paragraph()
ftr.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = ftr.add_run("MasPart · Part Number Finder")
r.italic = True
r.font.size = Pt(9)
r.font.color.rgb = MP_INK_50

# ── Save ─────────────────────────────────────────────────────────────
doc.save(str(OUT))
print(f"OK -> {OUT}")
